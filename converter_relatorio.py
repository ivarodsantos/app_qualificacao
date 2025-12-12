"""
Script melhorado para converter relatório técnico de Markdown para Word (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file, docx_file):
    """Converte arquivo Markdown para Word (.docx)"""
    
    # Criar documento Word
    doc = Document()
    
    # Ler arquivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    in_table = False
    table_headers = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Blocos de código
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'List Bullet'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Cabeçalhos
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Tabelas
        elif '|' in line and line.count('|') >= 2:
            # Detectar início de tabela
            if not in_table:
                in_table = True
                # Extrair cabeçalhos
                table_headers = [h.strip() for h in line.split('|') if h.strip()]
                # Criar tabela
                current_table = doc.add_table(rows=1, cols=len(table_headers))
                current_table.style = 'Light Grid Accent 1'
                # Preencher cabeçalhos
                for idx, header in enumerate(table_headers):
                    cell = current_table.rows[0].cells[idx]
                    cell.text = header
                    # Negrito no cabeçalho
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            else:
                # Linha de separação (---|---|---)
                if '---' in line or '===' in line:
                    pass  # Ignorar linha de separação
                else:
                    # Linha de dados
                    row_data = [d.strip() for d in line.split('|') if d.strip()]
                    if row_data and len(row_data) == len(table_headers):
                        row_cells = current_table.add_row().cells
                        for idx, data in enumerate(row_data):
                            row_cells[idx].text = data
        else:
            in_table = False
            
            # Listas
            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            # Linhas vazias
            elif line.strip() == '':
                doc.add_paragraph()
            # Separadores
            elif line.strip() == '---':
                p = doc.add_paragraph()
                p.add_run('_' * 80)
            # Texto normal
            elif line.strip():
                # Processar formatação básica
                clean_line = line
                # Remover > [!NOTE] etc
                clean_line = re.sub(r'^>\s*\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', clean_line)
                clean_line = re.sub(r'^>\s*', '', clean_line)
                
                # Adicionar parágrafo
                p = doc.add_paragraph()
                
                # Processar negrito e código inline
                parts = re.split(r'(\*\*.*?\*\*|`.*?`)', clean_line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        p.add_run(part)
        
        i += 1
    
    # Salvar documento
    doc.save(docx_file)
    return True

try:
    print('Convertendo relatorio tecnico para Word...')
    markdown_to_docx(
        'docs/relatorio_tecnico_filtro_projeto.md',
        'docs/relatorio_tecnico_filtro_projeto.docx'
    )
    print('[OK] Conversao concluida com sucesso!')
    print('  Arquivo criado: docs/relatorio_tecnico_filtro_projeto.docx')
    
except Exception as e:
    print(f'Erro ao converter: {e}')
    import traceback
    traceback.print_exc()
